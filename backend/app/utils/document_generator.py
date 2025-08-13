from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, Any, Optional
import os
from datetime import datetime
from app.utils.config import settings
import logging

logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self):
        self.template_dir = settings.template_dir
        self.deliverable_dir = settings.deliverable_dir
        
        # Ensure directories exist
        os.makedirs(self.deliverable_dir, exist_ok=True)
        os.makedirs(self.template_dir, exist_ok=True)
    
    def create_deliverable(
        self,
        client_name: str,
        deliverable_type: str,
        content_sections: Dict[str, str],
        stakeholder_name: Optional[str] = None
    ) -> str:
        """
        Create a MS Word deliverable document
        """
        try:
            # Create new document
            doc = Document()
            
            # Set up document styles
            self._setup_document_styles(doc)
            
            # Add header
            self._add_header(doc, client_name, deliverable_type, stakeholder_name)
            
            # Add content sections
            self._add_content_sections(doc, content_sections)
            
            # Add footer
            self._add_footer(doc)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{client_name.replace(' ', '_')}_{deliverable_type.replace(' ', '_')}_{timestamp}.docx"
            filepath = os.path.join(self.deliverable_dir, filename)
            
            # Save document
            doc.save(filepath)
            
            logger.info(f"Generated deliverable: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Failed to create deliverable: {e}")
            raise
    
    def _setup_document_styles(self, doc: Document):
        """Set up document styles"""
        # Title style
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(18)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1 style
        h1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Calibri'
        h1_font.size = Pt(14)
        h1_font.bold = True
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
        
        # Heading 2 style
        h2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Calibri'
        h2_font.size = Pt(12)
        h2_font.bold = True
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(6)
        
        # Body style
        body_style = doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15
    
    def _add_header(self, doc: Document, client_name: str, deliverable_type: str, stakeholder_name: Optional[str] = None):
        """Add document header"""
        # Title
        title = doc.add_paragraph(f"{deliverable_type.title()}", style='CustomTitle')
        
        # Subtitle
        subtitle_text = f"Prepared for {client_name}"
        if stakeholder_name:
            subtitle_text += f" - {stakeholder_name}"
        subtitle = doc.add_paragraph(subtitle_text, style='CustomBody')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", style='CustomBody')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_content_sections(self, doc: Document, content_sections: Dict[str, str]):
        """Add content sections to document"""
        section_mapping = {
            'executive_summary': 'Executive Summary',
            'key_recommendations': 'Key Recommendations',
            'background': 'Background and Context',
            'analysis': 'Analysis and Findings',
            'next_steps': 'Next Steps and Implementation Plan'
        }
        
        for section_key, content in content_sections.items():
            if content and content.strip():
                # Add section heading
                heading_text = section_mapping.get(section_key, section_key.replace('_', ' ').title())
                heading = doc.add_paragraph(heading_text, style='CustomHeading1')
                
                # Add content
                content_para = doc.add_paragraph(content, style='CustomBody')
                
                # Add spacing
                doc.add_paragraph()
    
    def _add_footer(self, doc: Document):
        """Add document footer"""
        # Add page break before footer
        doc.add_page_break()
        
        # Footer content
        footer_para = doc.add_paragraph("Jacob Meadow Associates", style='CustomBody')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_para2 = doc.add_paragraph("Confidential - For Internal Use Only", style='CustomBody')
        footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def create_template(self, template_name: str = "default_template.docx") -> str:
        """
        Create a default template for deliverables
        """
        try:
            doc = Document()
            
            # Set up styles
            self._setup_document_styles(doc)
            
            # Add template placeholders
            title = doc.add_paragraph("{{DELIVERABLE_TITLE}}", style='CustomTitle')
            
            subtitle = doc.add_paragraph("Prepared for {{CLIENT_NAME}}", style='CustomBody')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_para = doc.add_paragraph("Generated on {{DATE}}", style='CustomBody')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Add section placeholders
            sections = [
                "Executive Summary",
                "Key Recommendations", 
                "Background and Context",
                "Analysis and Findings",
                "Next Steps and Implementation Plan"
            ]
            
            for section in sections:
                heading = doc.add_paragraph(section, style='CustomHeading1')
                content = doc.add_paragraph("{{" + section.upper().replace(" ", "_") + "_CONTENT}}", style='CustomBody')
                doc.add_paragraph()
            
            # Save template
            template_path = os.path.join(self.template_dir, template_name)
            doc.save(template_path)
            
            logger.info(f"Created template: {template_path}")
            return template_path
            
        except Exception as e:
            logger.error(f"Failed to create template: {e}")
            raise

# Global document generator instance
document_generator = DocumentGenerator()