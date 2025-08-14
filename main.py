"""
main.py — JMA RAG Prototype (PostgreSQL + pgvector + Llama 3.2 Instruct + python-docx)

What this script does:
1) Connects to PostgreSQL (expects pgvector extension enabled).
2) (Optional) Loads schema/data if you call initialise_database(schema_path).
3) Embeds Knowledge_Entries and stores embeddings (vector(768)).
4) For each Client: retrieves top-K semantically similar snippets + stakeholder facts.
5) Builds a strict JSON prompt (persona, context, tone, guardrails).
6) Calls Llama 3.2 Instruct via Hugging Face Transformers to produce JSON.
7) Assembles a structured .docx (cover, stakeholders, Exec Summary, DAAEG, KPIs, Risks, Sources, Enrichment).
8) Saves the JSON “gold copy” back into Knowledge_Entries with an embedding for retrieval.

Environment variables you may set:
  PGHOST, PGPORT, PGUSER, PGPASSWORD, PGDATABASE
  HF_HOME / (run `huggingface-cli login` once to cache credentials)

Requirements are listed in requirements.txt.
"""

import os
import json
import hashlib
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Tuple

import psycopg2
import psycopg2.extras
from pgvector.psycopg2 import register_vector

from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline

from docx import Document  # python-docx

# -----------------------------
# Configuration & Connections
# -----------------------------

def get_db_connection():
    conn = psycopg2.connect(
        dbname=os.getenv("PGDATABASE", "jma_knowledge_base"),
        user=os.getenv("PGUSER", "postgres"),
        password=os.getenv("PGPASSWORD", ""),
        host=os.getenv("PGHOST", "localhost"),
        port=os.getenv("PGPORT", "5432"),
    )
    register_vector(conn)
    return conn

def ensure_output_dir():
    out = Path("deliverables")
    out.mkdir(parents=True, exist_ok=True)
    return out

# -----------------------------
# Optional: Initialize Schema
# -----------------------------

def initialise_database(schema_path: str):
    with open(schema_path, "r", encoding="utf-8") as f:
        sql = f.read()
    conn = get_db_connection()
    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql)
    finally:
        conn.close()

# -----------------------------
# Embeddings
# -----------------------------

def fetch_knowledge_needing_embeddings(conn) -> List[dict]:
    with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
        cur.execute("""
            SELECT id, content FROM knowledge_entries
            WHERE embedding IS NULL
        """)
        return [dict(r) for r in cur.fetchall()]

def store_embedding(conn, entry_id: int, emb: List[float]):
    with conn.cursor() as cur:
        cur.execute(
            "UPDATE knowledge_entries SET embedding = %s WHERE id = %s",
            (emb, entry_id)
        )

def compute_and_store_embeddings(embedder: SentenceTransformer):
    conn = get_db_connection()
    try:
        rows = fetch_knowledge_needing_embeddings(conn)
        if not rows:
            return
        texts = [r["content"] for r in rows]
        vecs = embedder.encode(texts, normalize_embeddings=True)
        with conn:
            for r, v in zip(rows, vecs):
                store_embedding(conn, r["id"], v.tolist())
    finally:
        conn.close()

# -----------------------------
# Retrieval
# -----------------------------

def load_client_and_stakeholders(conn, client_id: int) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
        cur.execute("SELECT id, name FROM clients WHERE id = %s", (client_id,))
        client = cur.fetchone()
        if not client:
            raise ValueError(f"Client {client_id} not found")
        cur.execute("""
            SELECT id, name, role, tone, priority1, priority2, priority3
            FROM stakeholders WHERE client_id = %s
        """, (client_id,))
        st = cur.fetchall()
    client_d = {"id": client["id"], "name": client["name"]}
    stakeholders = []
    for s in st:
        stakeholders.append({
            "id": s["id"],
            "name": s["name"],
            "role": s["role"],
            "tone": s["tone"],
            "priorities": [p for p in [s["priority1"], s["priority2"], s["priority3"]] if p]
        })
    return client_d, stakeholders

def semantic_retrieve(conn, client_id: int, embedder: SentenceTransformer, query_text: str, top_k: int = 5) -> List[Dict[str, Any]]:
    qvec = embedder.encode([query_text], normalize_embeddings=True)[0].tolist()
    with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
        cur.execute("""
            SELECT id, type, content
            FROM knowledge_entries
            WHERE client_id = %s
            ORDER BY embedding <-> %s
            LIMIT %s
        """, (client_id, qvec, top_k))
        return [dict(r) for r in cur.fetchall()]

def build_query_from_stakeholders(stakeholders: List[Dict[str, Any]]) -> str:
    # Simple heuristic query: concatenate priorities and tone words
    topics = []
    for s in stakeholders:
        topics.extend(s.get("priorities", []))
        if s.get("tone"):
            topics.append(f"tone:{s['tone']}")
    return " ; ".join(topics) or "project priorities; risks; timeline; cost; architecture"

# -----------------------------
# Prompt Engineering (Strict JSON)
# -----------------------------

def build_prompt(client: Dict[str, Any],
                 stakeholders: List[Dict[str, Any]],
                 retrieved_snippets: List[Dict[str, Any]]) -> str:
    # Persona + guardrails + structure
    today = datetime.now().date().isoformat()
    st_lines = "\n".join(
        f"- {s['name']} ({s['role']}), tone={s['tone']}, priorities={', '.join(s['priorities'])}"
        for s in stakeholders
    ) or "- (no stakeholders found)"

    ctx_lines = "\n".join(
        f"[KE-{r['id']}] {r['content'][:500].replace(chr(10),' ')}"
        for r in retrieved_snippets
    ) or "(no context retrieved)"

    prompt = f"""
You are an expert consultant from Jacob Meadow Associates. Be objective and fact-based. Do NOT invent facts.
Only use the provided context and stakeholder data. Avoid loaded or stereotypical language.

Return STRICT JSON with keys:
- cover: {{title, client, engagement, prepared_for, prepared_by, date, confidentiality}}
- stakeholders: [{{name, role, tone, priorities: [..]}}]
- executive_summary: {{paragraphs: [..], bullets: [..]}}
- current_situation: {{bullets: [..], sources: [..]}}
- recommendations: {{
    discover: [..], analyze: [..], architect: [..], execute: [..], govern: [..]
  }}
- kpis: [..]
- risks: [..]
- sources: [..]
- enrichment: {{gold_copy: true, knowledge_entry_id: "", model: "", prompt_hash: "", retrieval_ids: []}}

Constraints:
- Executive summary = 2 short paragraphs + 3 bullets (max).
- Use DAAEG headings exactly for recommendations.
- Include at least 3 KPIs and 3 risks.
- Use KE-style tags like [KE-1024] in 'sources' where appropriate.
- Ensure tone consistency per stakeholder (direct vs collaborative) in phrasing.

Client: {client['name']}
Stakeholders:
{st_lines}

Context (verbatim snippets for grounding):
{ctx_lines}

Cover defaults (you may override):
- title: "Client Readout – Executive Summary & Recommendations"
- engagement: "RAG Knowledge Base Prototype"
- prepared_for: "Executive Sponsor"
- prepared_by: "Jacob Meadow Associates"
- date: "{today}"
- confidentiality: "Confidential – For Client Use Only"

Output: JSON ONLY.
""".strip()
    return prompt

# -----------------------------
# LLM Call (Llama 3.2 Instruct)
# -----------------------------

def load_llm():
    # Make sure you've run: huggingface-cli login
    model_id = "meta-llama/Llama-3.2-3B-Instruct"
    tok = AutoTokenizer.from_pretrained(model_id)
    mdl = AutoModelForCausalLM.from_pretrained(model_id)
    gen = pipeline("text-generation", model=mdl, tokenizer=tok, max_new_tokens=900, temperature=0.2)
    return gen, model_id

def generate_report_with_llm(gen, prompt: str) -> str:
    out = gen(prompt)[0]["generated_text"]
    # Extract JSON from model output robustly (find first '{' to last '}')
    start = out.find("{")
    end = out.rfind("}")
    if start == -1 or end == -1 or end <= start:
        # Fallback minimal JSON structure if parsing fails
        return json.dumps({
            "cover": {"title": "Client Readout – Executive Summary & Recommendations",
                      "client": "UNKNOWN", "engagement": "RAG Knowledge Base Prototype",
                      "prepared_for": "Executive Sponsor", "prepared_by": "Jacob Meadow Associates",
                      "date": datetime.now().date().isoformat(),
                      "confidentiality": "Confidential – For Client Use Only"},
            "stakeholders": [], "executive_summary": {"paragraphs": [], "bullets": []},
            "current_situation": {"bullets": [], "sources": []},
            "recommendations": {"discover": [], "analyze": [], "architect": [], "execute": [], "govern": []},
            "kpis": [], "risks": [], "sources": [], "enrichment": {"gold_copy": True, "knowledge_entry_id": "", "model": "", "prompt_hash": "", "retrieval_ids": []}
        }, ensure_ascii=False)
    js = out[start:end+1]
    # Ensure it's valid JSON
    try:
        json.loads(js)
        return js
    except json.JSONDecodeError:
        # Try to fix trailing commas and retry (very light-touch)
        js2 = js.replace(",]", "]").replace(",}", "}")
        json.loads(js2)
        return js2

# -----------------------------
# DOCX Assembly (python-docx)
# -----------------------------

def create_word_document_from_json(file_path: Path, json_text: str):
    data = json.loads(json_text)
    doc = Document()

    # Cover / Title block
    c = data.get("cover", {})
    doc.add_heading(c.get("title", "Client Readout – Executive Summary & Recommendations"), level=0)
    for line in [
        f"Client: {c.get('client','')}",
        f"Engagement: {c.get('engagement','')}",
        f"Prepared for: {c.get('prepared_for','')}",
        f"Prepared by: {c.get('prepared_by','')}",
        f"Date: {c.get('date','')}",
        c.get("confidentiality","Confidential – For Client Use Only"),
    ]:
        if line.strip():
            doc.add_paragraph(line)

    # Stakeholders
    doc.add_heading("Key Stakeholders", level=1)
    st = data.get("stakeholders", [])
    if st:
        tbl = doc.add_table(rows=1, cols=4)
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Name", "Role", "Tone", "Top Priorities"
        for s in st:
            row = tbl.add_row().cells
            row[0].text = s.get("name","")
            row[1].text = s.get("role","")
            row[2].text = s.get("tone","")
            row[3].text = "; ".join(s.get("priorities", []))
    else:
        doc.add_paragraph("No stakeholders found.")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    es = data.get("executive_summary", {})
    for para in es.get("paragraphs", []):
        doc.add_paragraph(para)
    for b in es.get("bullets", []):
        doc.add_paragraph(b, style="List Bullet")

    # Current Situation
    doc.add_heading("Current Situation Assessment", level=1)
    cs = data.get("current_situation", {})
    for b in cs.get("bullets", []):
        doc.add_paragraph(b, style="List Bullet")
    srcs = cs.get("sources", [])
    if srcs:
        doc.add_paragraph("Sources referenced in assessment: " + ", ".join(srcs))

    # Recommendations (DAAEG)
    doc.add_heading("Key Recommendations (DAAEG)", level=1)
    recs = data.get("recommendations", {})
    for sec in ["discover","analyze","architect","execute","govern"]:
        doc.add_heading(sec.capitalize(), level=2)
        for b in recs.get(sec, []):
            doc.add_paragraph(b, style="List Bullet")

    # KPIs
    doc.add_heading("Expected Outcomes & KPIs", level=1)
    for k in data.get("kpis", []):
        doc.add_paragraph(k, style="List Bullet")

    # Risks
    doc.add_heading("Risks & Mitigations", level=1)
    for r in data.get("risks", []):
        doc.add_paragraph(r, style="List Bullet")

    # Sources & Traceability
    doc.add_heading("Sources & Traceability", level=1)
    for s in data.get("sources", []):
        doc.add_paragraph(s, style="List Bullet")

    # Enrichment Metadata (Closed-Loop)
    doc.add_heading("Enrichment Metadata (Closed-Loop)", level=1)
    e = data.get("enrichment", {})
    for line in [
        f"GoldCopy: {e.get('gold_copy', False)}",
        f"KnowledgeEntryId: {e.get('knowledge_entry_id','')}",
        f"Model: {e.get('model','')}",
        f"Prompt Hash: {e.get('prompt_hash','')}",
        "Retrieval Context IDs: " + ", ".join(e.get("retrieval_ids", [])),
    ]:
        doc.add_paragraph(line)

    doc.save(str(file_path))

# -----------------------------
# Closed-loop: store gold copy
# -----------------------------

def insert_deliverable_json(client_id: int, json_text: str, embedder: SentenceTransformer):
    conn = get_db_connection()
    try:
        vec = embedder.encode([json_text], normalize_embeddings=True)[0].tolist()
        with conn:
            with conn.cursor() as cur:
                cur.execute("""
                    INSERT INTO knowledge_entries (client_id, stakeholder_id, type, content, embedding)
                    VALUES (%s, NULL, 'deliverable', %s, %s)
                """, (client_id, json_text, vec))
    finally:
        conn.close()

# -----------------------------
# Main pipeline
# -----------------------------

def run():
    ensure_output_dir()

    # Load embedder & LLM
    embedder = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
    gen, model_id = load_llm()

    # Compute embeddings for any rows that are missing them
    compute_and_store_embeddings(embedder)

    # Iterate over clients
    conn = get_db_connection()
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cur:
            cur.execute("SELECT id, name FROM clients ORDER BY id")
            clients = cur.fetchall()
    finally:
        conn.close()

    for c in clients:
        client_id = c["id"]
        client_name = c["name"]

        conn = get_db_connection()
        try:
            client_d, stakeholders = load_client_and_stakeholders(conn, client_id)
            # Build a retrieval query from stakeholder priorities
            q = build_query_from_stakeholders(stakeholders)
            # Hybrid-ish: just a semantic retrieve using the query
            retrieved = semantic_retrieve(conn, client_id, embedder, q, top_k=5)
        finally:
            conn.close()

        # Prompt
        prompt = build_prompt(client_d, stakeholders, retrieved)

        # LLM generate (JSON)
        report_json = generate_report_with_llm(gen, prompt)

        # Add model identity & prompt hash in the JSON payload (light touch)
        try:
            data = json.loads(report_json)
        except Exception:
            data = {}
        e = data.get("enrichment", {})
        e["model"] = e.get("model") or model_id
        e["prompt_hash"] = e.get("prompt_hash") or hashlib.sha1(prompt.encode("utf-8")).hexdigest()[:12]
        e["retrieval_ids"] = e.get("retrieval_ids") or [f"KE-{r['id']}" for r in retrieved]
        data["enrichment"] = e
        # also ensure cover.client is set
        cov = data.get("cover", {})
        if not cov.get("client"):
            cov["client"] = client_name
        data["cover"] = cov
        report_json = json.dumps(data, ensure_ascii=False)

        # DOCX
        out_path = ensure_output_dir() / f"deliverable_{client_name.replace(' ', '_')}.docx"
        create_word_document_from_json(out_path, report_json)

        # Closed-loop store deliverable
        insert_deliverable_json(client_id, report_json, embedder)

        print(f"✓ Generated deliverable for {client_name}: {out_path}")

    print("Pipeline complete.")

if __name__ == "__main__":
    run()
